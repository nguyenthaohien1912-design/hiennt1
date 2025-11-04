# python.py

import streamlit as st
from docx import Document
import unicodedata, re, os
from difflib import get_close_matches, SequenceMatcher

# =======================
# Cấu hình trang
# =======================
st.set_page_config(
    page_title="Sổ tay hướng dẫn kiểm tra Agribank Hà Thành",
    page_icon="📘",
    layout="wide"
)

# =======================
# CSS chủ đề Agribank
# =======================
st.markdown("""
<style>
.main { background:#fff; color:#222; font-family:'Segoe UI', sans-serif; }
.block-container { padding-top:1rem; padding-bottom:1rem; }
h1,h2,h3,h4 { color:#800000 !important; }
.sidebar .sidebar-content { background:#8B0000; color:#fff; }
.stButton button { background:#800000; color:#fff; border-radius:6px; border:none; padding:.5rem 1rem; }
.stButton button:hover { background:#a00000; color:#fff; }
</style>
""", unsafe_allow_html=True)

# =======================
# Logo + tiêu đề
# =======================
col1, col2 = st.columns([0.15, 0.85])
with col1:
    # dùng logo cục bộ nếu có, fallback sang online để tránh lỗi
    logo_path = "logo_agribank.png"
    if os.path.exists(logo_path):
        st.image(logo_path, use_column_width=True)
    else:
        st.image("https://upload.wikimedia.org/wikipedia/commons/4/4b/Agribank_logo.png",
                 use_column_width=True)
with col2:
    st.title("📘 SỔ TAY HƯỚNG DẪN KIỂM TRA NGHIỆP VỤ")
    st.subheader("Agribank Chi nhánh Hà Thành – Phiên bản số hóa")
st.markdown("---")

# =======================
# Tiện ích văn bản
# =======================
def normalize_text(text: str) -> str:
    """Bỏ dấu tiếng Việt + đưa về lower để so khớp không dấu."""
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c)).lower()

@st.cache_data
def load_docx(file_path):
    from docx import Document
    import os
    if not os.path.exists(file_path):
        st.error(f"❌ Không tìm thấy file: {file_path}")
        st.stop()

    doc = Document(file_path)
    chapters = {}
    current_chapter = "Khác"

    def extract_text_from_table(table):
        """Đọc toàn bộ nội dung từ bảng và nối lại thành các dòng văn bản"""
        rows = []
        for row in table.rows:
            # Lấy text từ từng cell trong bảng
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        return rows

    # Đọc toàn bộ phần thân tài liệu (paragraphs + tables)
    for block in doc.element.body:
        # Đoạn văn
        if block.tag.endswith('p'):
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if text.lower().startswith("chương"):
                    current_chapter = text
                    chapters[current_chapter] = []
                else:
                    chapters.setdefault(current_chapter, []).append(text)
            break  # tránh đọc lại các đoạn trùng
        # Bảng
        elif block.tag.endswith('tbl'):
            # tìm tất cả bảng
            for table in doc.tables:
                for t in extract_text_from_table(table):
                    chapters.setdefault(current_chapter, []).append(t)
            break

    return chapters
    
# =======================
# Nạp dữ liệu
# =======================
FILENAME = "So_tay_Agribank.docx.docx"   # đổi tên file tại đây nếu cần
chapters = load_docx(FILENAME)

# Chuẩn bị corpus phẳng để fuzzy/gợi ý
@st.cache_data
def build_corpus(chapters_dict):
    rows = []
    for ch, paras in chapters_dict.items():
        for p in paras:
            rows.append({
                "chapter": ch,
                "text": p,
                "norm": normalize_text(p)
            })
    return rows

corpus = build_corpus(chapters)

# Từ điển đồng nghĩa/biến thể hay dùng (có thể bổ sung dần)
SYNONYMS = {
    "ho so cap tin dung": [
        "hồ sơ cấp tín dụng", "hồ sơ vay vốn", "bộ hồ sơ tín dụng",
        "hồ sơ cho vay", "hồ sơ khoản vay", "hồ sơ tín dụng"
    ],
    "tin dung": ["tín dụng", "cho vay", "khoản vay", "cấp tín dụng"],
    "bao dam": ["bảo đảm", "tài sản bảo đảm", "tsbđ", "thế chấp", "cầm cố"],
    "thanh toan": ["thanh toán", "kế toán", "chứng từ", "hạch toán"],
}

def expand_query(q: str):
    base = normalize_text(q)
    variants = {q}  # giữ nguyên bản có dấu
    variants.add(base)  # bản không dấu
    for key, alts in SYNONYMS.items():
        if key in base:
            variants.update(alts)
            variants.update([normalize_text(a) for a in alts])
    return list({v for v in variants if v})

def highlight(text: str, variants):
    """Tô đậm tất cả biến thể (có dấu & không dấu)."""
    # pattern ghép OR cho các biến thể có dấu
    with_diacritics = [v for v in variants if any("ăâđêôơưáàảãạéèẻẽẹóòỏõọúùủũụíìỉĩịýỳỷỹỵ" in v.lower() for _ in [0])]
    # tô đậm bản có dấu trước:
    if with_diacritics:
        pattern = r"(" + "|".join(re.escape(v) for v in sorted(with_diacritics, key=len, reverse=True)) + r")"
        text = re.sub(pattern, r"**\1**", text, flags=re.IGNORECASE)
    # nếu không có dấu trong biến thể, bỏ qua highlight không dấu để tránh bôi đậm sai vị trí
    return text

def search_documents(query: str, scope_dict, use_synonyms=True):
    """Tìm trong scope_dict: {chapter:[paras]} -> trả về {chapter:[paras match]}"""
    variants = expand_query(query) if use_synonyms else [query, normalize_text(query)]
    norm_variants = [normalize_text(v) for v in variants]
    results = {}
    for ch, paras in scope_dict.items():
        hits = []
        for p in paras:
            pn = normalize_text(p)
            if any(v in pn for v in norm_variants):
                hits.append(p)
        if hits:
            results[ch] = hits
    return results, variants

# =======================
# Sidebar
# =======================
sb_logo = "logo_agribank.png"
if os.path.exists(sb_logo):
    st.sidebar.image(sb_logo, use_column_width=True)
else:
    st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/4/4b/Agribank_logo.png",
                     use_column_width=True)

st.sidebar.markdown("### 📑 **Danh mục chương**")
chapter_list = list(chapters.keys())
selected_chapter = st.sidebar.selectbox("Chọn chương:", chapter_list)

st.sidebar.markdown("---")
st.sidebar.markdown("### 💬 **Chatbot hướng dẫn kiểm tra**")
query = st.sidebar.text_input("Nhập từ khóa/câu hỏi (VD: hồ sơ cấp tín dụng, chứng từ...)")
search_in_current = st.sidebar.checkbox("🔍 Chỉ tìm trong chương đã chọn", value=False)
use_syn = st.sidebar.checkbox("✨ Mở rộng từ đồng nghĩa", value=True)

st.sidebar.markdown("---")
with open(FILENAME, "rb") as f:
    st.sidebar.download_button(
        label="⬇️ Tải Sổ tay gốc (.docx)",
        data=f,
        file_name="So_tay_Agribank.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =======================
# Nội dung chương đã chọn
# =======================
st.header(f"📂 {selected_chapter}")
for para in chapters[selected_chapter]:
    st.markdown(f"- {para}")

# =======================
# Tìm kiếm nâng cao + gợi ý
# =======================
if query:
    st.markdown("---")
    st.subheader(f"🔎 Kết quả tìm kiếm cho: *{query}*")

    scope = {selected_chapter: chapters[selected_chapter]} if search_in_current else chapters
    results_by_chapter, variants = search_documents(query, scope, use_synonyms=use_syn)

    if results_by_chapter:
        for ch, paras in results_by_chapter.items():
            with st.expander(f"📁 {ch} ({len(paras)} kết quả)", expanded=True):
                for p in paras:
                    st.markdown(f"🔹 {highlight(p, variants)}")
    else:
        st.info("Không tìm thấy nội dung khớp hoàn toàn. Dưới đây là các **gợi ý gần đúng**:")
        # Gợi ý gần đúng từ toàn bộ corpus đã chuẩn hóa
        full_texts = [row["text"] for row in corpus]
        # Lấy 10 câu/đoạn giống nhất theo ratio
        scored = sorted(
            [(t, SequenceMatcher(None, normalize_text(query), normalize_text(t)).ratio()) for t in full_texts],
            key=lambda x: x[1],
            reverse=True
        )[:10]
        for t, score in scored:
            # chỉ hiển thị gợi ý đủ “giống”
            if score >= 0.45:
                st.markdown(f"💡 {t}  \n&emsp;`similarity: {score:.2f}`")
        if not scored or scored[0][1] < 0.45:
            st.write("• Thử rút gọn từ khóa (vd: *hồ sơ*, *tín dụng*, *cho vay*, *chứng từ*).")
